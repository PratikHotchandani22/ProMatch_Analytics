from fpdf import FPDF
from docx import Document
import streamlit as st
from io import BytesIO
import tiktoken

# Pricing dictionary for different models
# Updated Pricing dictionary for different models
PRICING = {
    "chatgpt-4o-latest": {"input": 5.00 / 1_000_000, "output": 15.00 / 1_000_000},
    "gpt-4-turbo": {"input": 10.00 / 1_000_000, "output": 30.00 / 1_000_000},
    "gpt-4-turbo-2024-04-09": {"input": 10.00 / 1_000_000, "output": 30.00 / 1_000_000},
    "gpt-4": {"input": 30.00 / 1_000_000, "output": 60.00 / 1_000_000},
    "gpt-4-32k": {"input": 60.00 / 1_000_000, "output": 120.00 / 1_000_000},
    "gpt-4-0125-preview": {"input": 10.00 / 1_000_000, "output": 30.00 / 1_000_000},
    "gpt-4-1106-preview": {"input": 10.00 / 1_000_000, "output": 30.00 / 1_000_000},
    "gpt-4-vision-preview": {"input": 10.00 / 1_000_000, "output": 30.00 / 1_000_000},
    "gpt-3.5-turbo-0125": {"input": 0.50 / 1_000_000, "output": 1.50 / 1_000_000},
    "gpt-3.5-turbo-instruct": {"input": 1.50 / 1_000_000, "output": 2.00 / 1_000_000},
    "gpt-3.5-turbo-1106": {"input": 1.00 / 1_000_000, "output": 2.00 / 1_000_000},
    "gpt-3.5-turbo-0613": {"input": 1.50 / 1_000_000, "output": 2.00 / 1_000_000},
    "gpt-3.5-turbo-16k-0613": {"input": 3.00 / 1_000_000, "output": 4.00 / 1_000_000},
    "gpt-3.5-turbo-0301": {"input": 1.50 / 1_000_000, "output": 2.00 / 1_000_000},
    # Adding the new model
    "gpt-4o-mini": {"input": 0.150 / 1_000_000, "output": 0.600 / 1_000_000},  # Updated model pricing
}

# Manually define encodings for specific models
MODEL_ENCODINGS = {
    "chatgpt-4o-latest": "o200k_base",             # For gpt-4o, gpt-4o-mini
    "gpt-4-turbo": "cl100k_base",                 # For gpt-4-turbo, gpt-4, gpt-3.5-turbo
    "gpt-4-turbo-2024-04-09": "cl100k_base",      # For gpt-4-turbo, gpt-4, gpt-3.5-turbo
    "gpt-4": "cl100k_base",                       # For gpt-4-turbo, gpt-4, gpt-3.5-turbo
    "gpt-4-32k": "cl100k_base",                   # For gpt-4 (32k token limit)
    "gpt-4-0125-preview": "cl100k_base",          # For gpt-4-turbo, gpt-4, gpt-3.5-turbo
    "gpt-4-1106-preview": "cl100k_base",          # For gpt-4-turbo, gpt-4, gpt-3.5-turbo
    "gpt-4-vision-preview": "cl100k_base",        # For gpt-4-turbo, gpt-4, gpt-3.5-turbo
    "gpt-3.5-turbo-0125": "cl100k_base",          # For gpt-3.5-turbo, gpt-4-turbo, gpt-4
    "gpt-3.5-turbo-instruct": "cl100k_base",      # For gpt-3.5-turbo-instruct, gpt-4, gpt-3.5-turbo
    "gpt-3.5-turbo-1106": "cl100k_base",          # For gpt-3.5-turbo, gpt-4, gpt-4-turbo
    "gpt-3.5-turbo-0613": "cl100k_base",          # For gpt-3.5-turbo, gpt-4, gpt-4-turbo
    "gpt-3.5-turbo-16k-0613": "cl100k_base",      # For gpt-3.5-turbo, gpt-4, gpt-4-turbo
    "gpt-3.5-turbo-0301": "cl100k_base",   
    "gpt-4o-mini": "o200k_base",       # For gpt-3.5-turbo, gpt-4, gpt-4-turbo
}


def calculate_token_cost(input_text, output_text, budget=10.00, model_name=None):
    """
    Calculate the token count and cost for processing input and output text using specified OpenAI models,
    and determine how many times the code can be run with a given budget.

    Args:
        input_text (str): The input text to be processed.
        output_text (str): The generated output text.
        budget (float): The total budget available for running the code (default is 10 dollars).
        model_name (str or list): The model name or a list of model names to include in the calculation. 
                                  If None, all available models are included.

    Returns:
        list: A list of dictionaries containing token and cost information for each model,
              along with how many times the code can be run within the budget.
    """
    results = []

    # Determine which models to process
    models_to_process = (
        [model_name] if isinstance(model_name, str) else model_name
        if isinstance(model_name, list) else PRICING.keys()
    )

    # Loop through the selected models and calculate costs
    for model in models_to_process:
        # Ensure the model exists in the pricing dictionary
        if model not in PRICING:
            print(f"Warning: Model {model} is not available in PRICING. Skipping.")
            continue

        # Get the correct tokenizer encoding for the model
        encoding_name = MODEL_ENCODINGS.get(model, None)
        if encoding_name is None:
            print(f"Warning: No encoding for {model}. Skipping.")
            continue  # Skip models that don't have a valid encoding

        # Get the encoder for the model
        encoder = tiktoken.get_encoding(encoding_name)

        # Calculate token counts
        input_tokens = len(encoder.encode(input_text))
        output_tokens = len(encoder.encode(output_text))
        total_tokens = input_tokens + output_tokens

        # Calculate costs
        pricing = PRICING[model]
        input_cost = input_tokens * pricing["input"]
        output_cost = output_tokens * pricing["output"]
        total_cost = input_cost + output_cost

        # Calculate how many runs can be done with the given budget
        runs_possible = budget / total_cost if total_cost > 0 else 0

        # Store result for this model
        results.append({
            "model_name": model,
            "input_tokens": input_tokens,
            "output_tokens": output_tokens,
            "total_tokens": total_tokens,
            "total_cost": total_cost,
            "runs_possible_with_budget": runs_possible
        })

    return results


def calculate_token_cost_for_all_models(input_text, output_text, budget=10.00):
    """
    Calculate the token count and cost for processing input and output text using all available OpenAI models,
    and determine how many times the code can be run with a given budget.

    Args:
        input_text (str): The input text to be processed.
        output_text (str): The generated output text.
        budget (float): The total budget available for running the code (default is 10 dollars).

    Returns:
        list: A list of dictionaries containing token and cost information for each model,
              along with how many times the code can be run within the budget.
    """
    results = []

    # Loop through all models and calculate costs
    for model_name, pricing in PRICING.items():
        # Get the correct tokenizer encoding for the model
        encoding_name = MODEL_ENCODINGS.get(model_name, None)
        if encoding_name is None:
            print(f"Warning: No encoding for {model_name}. Skipping.")
            continue  # Skip models that don't have a valid encoding

        # Get the encoder for the model
        encoder = tiktoken.get_encoding(encoding_name)

        # Calculate token counts
        input_tokens = len(encoder.encode(input_text))
        output_tokens = len(encoder.encode(output_text))
        total_tokens = input_tokens + output_tokens

        # Calculate costs
        input_cost = input_tokens * pricing["input"]
        output_cost = output_tokens * pricing["output"]
        total_cost = input_cost + output_cost

        # Calculate how many runs can be done with the given budget
        runs_possible = budget / total_cost if total_cost > 0 else 0

        # Store result for this model
        results.append({
            "model_name": model_name,
            "input_tokens": input_tokens,
            "output_tokens": output_tokens,
            "total_tokens": total_tokens,
            "total_cost": total_cost,
            "runs_possible_with_budget": runs_possible
        })

    return results

# Function to save cover letter as PDF
def save_as_pdf(content):
    # Initialize PDF object
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    # Split content into lines for writing
    lines = content.split("\n")
    for line in lines:
        pdf.cell(0, 10, line, ln=True)

    # Output PDF to a string and write to BytesIO buffer
    pdf_buffer = BytesIO()
    pdf_data = pdf.output(dest="S").encode("latin1")  # Output as string and encode to bytes
    pdf_buffer.write(pdf_data)
    pdf_buffer.seek(0)

    return pdf_buffer.getvalue()  # Return the data as bytes

# Function to save cover letter as Word document
def save_as_docx(content):
    doc = Document()
    doc.add_heading("Cover Letter", level=1)

    # Add the cover letter content
    for line in content.split("\n"):
        doc.add_paragraph(line)

    # Save to a BytesIO object
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)

    return docx_buffer.getvalue()  # Return the data as bytes